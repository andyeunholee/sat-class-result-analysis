#!/usr/bin/env python3
"""
Elite Prep - SAT Class Test Result Analysis for Teachers
=========================================================
Give it a folder of student SAT/DSAT practice-test score-report PDFs
(one class, one test) and it generates a Word report in the exact format
of the "Sample-DSAT-06-A Result Analysis Teacher Report" reference document:

    "<TEST CODE> SAT Test Result Analysis for Teacher.docx"

For each of the four test sections (Math Section 1 & 2, English Section
1 & 2) the questions are listed from most-missed to least-missed.

Rules always enforced:
  * The report is written in English.
  * Branding is always just "Elite Prep" (never "Elite Prep Suwanee").
  * No "Andy Lee, Director, ..." signature line.
  * No footers on any page.
  * Test code and test date are auto-detected from the PDFs.

Usage:
    python generate_class_report.py <folder_with_pdfs>
        [--test-code DSAT-05-A] [--test-date "June 29, 2026"]
        [--output-dir DIR] [--dump-text]

--dump-text writes each PDF's raw extracted text to <name>.extracted.txt
so the parser can be adapted if a new score-report layout appears.
"""

import argparse
import datetime
import os
import re
import sys
from collections import defaultdict
from pathlib import Path

import pdfplumber

# Load ANTHROPIC_API_KEY from a .env file next to this script, if present.
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).resolve().parent / ".env")
except ImportError:
    pass
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

# ---------------------------------------------------------------------------
# Design system (measured directly from the reference document)
# ---------------------------------------------------------------------------

EN = "–"                       # – en dash
EM = "—"                       # — em dash
CIRCLED = {1: "①", 2: "②", 3: "③", 4: "④"}  # ① ② ③ ④

RW_DOMAINS = [
    "Craft and Structure",
    "Information and Ideas",
    "Standard English Conventions",
    "Expression of Ideas",
]
MATH_DOMAINS = [
    "Algebra",
    "Advanced Math",
    "Problem-Solving and Data Analysis",
    "Geometry and Trigonometry",
]

DOMAIN_ALIASES = {
    "craft and structure": "Craft and Structure",
    "information and ideas": "Information and Ideas",
    "standard english conventions": "Standard English Conventions",
    "expression of ideas": "Expression of Ideas",
    "algebra": "Algebra",
    "advanced math": "Advanced Math",
    "problem-solving and data analysis": "Problem-Solving and Data Analysis",
    "problem solving and data analysis": "Problem-Solving and Data Analysis",
    "problem-solving & data analysis": "Problem-Solving and Data Analysis",
    "geometry and trigonometry": "Geometry and Trigonometry",
    "geometry & trigonometry": "Geometry and Trigonometry",
}


def sanitize_branding(text: str) -> str:
    """Branding must always be written as just 'Elite Prep'."""
    return re.sub(r"Elite\s*Prep\s*Suwanee", "Elite Prep", text, flags=re.I)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

class QuestionResult:
    __slots__ = ("section", "module", "number", "status", "domain", "gridin")

    def __init__(self, section, module, number, status,
                 domain=None, gridin=None):
        self.section = section   # "RW" or "Math"
        self.module = module     # 1 or 2
        self.number = number     # question number (int)
        self.status = status     # "correct" | "incorrect" | "omitted"
        self.domain = domain     # canonical skill-area name or None
        self.gridin = gridin     # True if student-produced response, else None


class StudentResult:
    def __init__(self, source):
        self.source = source     # file name (never written into the report)
        self.student_name = None  # detected only so it can be scrubbed
        self.total = None        # scaled scores (int or None)
        self.rw = None
        self.math = None
        self.questions = []      # list[QuestionResult]
        self.test_code = None
        self.test_date = None
        self.domain_counts = {}  # {skill area: (correct, total)} from page 1

    @property
    def parsed_ok(self):
        return bool(self.questions)


# ---------------------------------------------------------------------------
# PDF text extraction
# ---------------------------------------------------------------------------

def extract_pdf_text(source) -> str:
    """Extract text from a PDF path or file-like object (e.g. an upload).
    Table rows often extract more cleanly than the raw text layer, so both
    are captured."""
    chunks = []
    with pdfplumber.open(source) as pdf:
        for page in pdf.pages:
            if chunks:
                chunks.append("\f")  # page break marker (multi-student files)
            chunks.append(page.extract_text() or "")
            try:
                for table in page.extract_tables():
                    for row in table:
                        cells = [c.strip() for c in row if c and c.strip()]
                        if cells:
                            chunks.append("  ".join(cells))
            except Exception:
                pass  # a malformed table never blocks the text layer
    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# Score-report parsing
# ---------------------------------------------------------------------------

STATUS_WORDS = {
    "correct": "correct",
    "incorrect": "incorrect",
    "omitted": "omitted",
    "unanswered": "omitted",
    "no answer": "omitted",
    "skipped": "omitted",
}

SECTION_PATTERNS = [
    (re.compile(r"reading\s*(?:and|&)\s*writing", re.I), "RW"),
    (re.compile(r"\bR\s*&\s*W\b", re.I), "RW"),
    (re.compile(r"\benglish\b", re.I), "RW"),
    (re.compile(r"\bmath\b", re.I), "Math"),
]

MODULE_PATTERN = re.compile(r"(?:module|section)\s*([12])\b", re.I)

# Row style A - section and module named on the row itself:
#   "12 Reading and Writing: Module 1 B A Incorrect"
ROW_A = re.compile(
    r"^\s*(?:Question\s*)?(\d{1,2})[.)]?\s+"
    r"(Reading\s*(?:and|&)\s*Writing|English|Math)\s*[:\-]?\s*"
    r"(?:Module|Section)\s*([12])\b(.*)$",
    re.I,
)

# Row style B - section/module known from a header line above:
#   "12 B A Incorrect" / "12 Craft and Structure B A" / "12 Incorrect"
ROW_B = re.compile(r"^\s*(?:Question\s*)?(\d{1,2})[.)]?\s+(\S.*)$")

SCORE_PATTERNS = {
    "total": re.compile(
        r"total\s*(?:sat\s*)?score\s*[:\-]?\s*(\d{3,4})", re.I),
    "rw": re.compile(
        r"reading\s*(?:and|&)\s*writing(?:\s*(?:section)?\s*score)?"
        r"\s*[:\-]?\s*(\d{3})\b", re.I),
    "math": re.compile(
        r"math(?:\s*(?:section)?\s*score)?\s*[:\-]?\s*(\d{3})\b", re.I),
}

TEST_CODE_PATTERNS = [
    re.compile(r"test\s*code\s*[:\-]?\s*"
               r"([A-Za-z0-9][A-Za-z0-9\-_ ]{1,30}?)\s*$", re.I | re.M),
    re.compile(r"\b(DSAT[-_ ]?\d+[-_ ]?[A-Z]?)\b", re.I),
    re.compile(r"\b(SAT\s*Practice\s*(?:Test\s*)?#?\s*\d+)\b", re.I),
    re.compile(r"\b(Practice\s*Test\s*\d+)\b", re.I),
    re.compile(r"\b(Bluebook\s*Practice\s*#?\s*\d+)\b", re.I),
]

DATE_PATTERNS = [
    re.compile(
        r"\b(January|February|March|April|May|June|July|August|September|"
        r"October|November|December)\s+\d{1,2},\s*\d{4}\b"),
    re.compile(r"\b\d{1,2}/\d{1,2}/\d{4}\b"),
]

# Student names are detected ONLY so the report can be checked/scrubbed.
NAME_PATTERNS = [
    re.compile(r"(?:student\s*name|name|student|test\s*taker)[ 	]*[:\-][ 	]*"
               r"([A-Za-z][A-Za-z'\-]+(?:[ 	]+[A-Za-z][A-Za-z'\-]+){0,3})",
               re.I),
]

MC_CHOICES = {"A", "B", "C", "D"}
OMIT_TOKENS = {"omitted", "unanswered", "--", "-", "blank", "—", "–"}


def eval_fraction(s: str):
    if re.fullmatch(r"-?\d+/\d+", s):
        num, den = s.split("/")
        return float(num) / float(den)
    return float(s)


def norm_answer(a: str) -> str:
    """Normalize an answer token so '3/4' == '.75' and 'b' == 'B'."""
    a = a.strip().strip(".,;").upper()
    try:
        return format(eval_fraction(a), ".6g")
    except Exception:
        return a


def extract_status_and_domain(fragment: str):
    """Pull correct/incorrect/omitted and a skill-area name out of a row
    fragment. Returns (status_or_None, domain_or_None, remainder)."""
    frag = fragment.strip()
    low = frag.lower()

    domain = None
    for alias, canonical in DOMAIN_ALIASES.items():
        if alias in low:
            domain = canonical
            idx = low.find(alias)
            frag = (frag[:idx] + frag[idx + len(alias):]).strip(" -:;|")
            low = frag.lower()
            break

    status = None
    for word, st in STATUS_WORDS.items():
        if re.search(r"\b" + re.escape(word) + r"\b", low):
            status = st
            break

    return status, domain, frag


def answers_from_fragment(fragment: str):
    """Best-effort (correct_answer, your_answer) from the tail of a row,
    e.g. 'B A', 'C Omitted', '3/4 .75', '104 --'. Returns (None, None) if
    the row doesn't look like an answer pair."""
    frag = fragment.strip(" .|")
    if not frag:
        return None, None
    tokens = frag.split()
    if len(tokens) < 2:
        return None, None
    return tokens[-2], tokens[-1]


def infer_status(correct_ans, your_ans):
    if not correct_ans or not your_ans:
        return None
    if your_ans.lower().strip(".,;") in OMIT_TOKENS:
        return "omitted"
    return ("correct" if norm_answer(correct_ans) == norm_answer(your_ans)
            else "incorrect")


def is_gridin_answer(correct_ans) -> bool:
    """A correct answer that is not a single A-D letter means a
    student-produced response (grid-in) question."""
    return bool(correct_ans) and norm_answer(correct_ans) not in MC_CHOICES


# ---------------------------------------------------------------------------
# Elite Prep "TEST SCORE REPORT" layout
# ---------------------------------------------------------------------------
# Page 1:  "ARNAV KUNDE • 113732102 AUG. 21, 2026 • DSK2402UU"
#          "Total\n1320\n400to1600", "Reading and Writing 570", "Math 750"
# Page 2:  four answer tables (R&W Module 1, R&W Module 2, Math Module 1,
#          Math Module 2).  Each row is "<q#> <correct answer> <response>"
#          where the response is "OK" (correct), "OMIT" (unanswered) or the
#          student's wrong answer.  pdfplumber emits the table titles out of
#          step with the rows, so tables are identified by row-block order.

# "<q#> <correct> <response>" with an optional 4th column: the question-type
# (skill-area) letter printed on newer reports, e.g. "19 C B S".
ELITE_ROW = re.compile(r"^\s*(\d{1,2})\s+(\S+)\s+(\S+)(?:\s+([A-Z]))?\s*$")
ELITE_SKILL_CODES = {
    "RW": {"C": "Craft and Structure", "I": "Information and Ideas",
           "S": "Standard English Conventions", "E": "Expression of Ideas"},
    "Math": {"A": "Algebra", "D": "Advanced Math",
             "P": "Problem-Solving and Data Analysis",
             "G": "Geometry and Trigonometry"},
}
# Page-1 per-skill summary, e.g. "CraftandStructure 86%\n12of14correct"
ELITE_SKILL_SUMMARY = re.compile(
    r"([A-Za-z][A-Za-z\-]+)\s*(\d{1,3})%\s*\n?\s*(\d+)\s*of\s*(\d+)\s*correct",
    re.I)
ELITE_TABLE_HEADER = re.compile(
    r"Reading\s*and\s*Writing\s*Mod-?\s*ule\s*([12])|Math\s*Module\s*([12])",
    re.I)
ELITE_ID_LINE = re.compile(
    r"^(?P<name>[A-Z][A-Za-z'\-]+(?:[ \t]+[A-Z][A-Za-z'\-]+){1,4})[ \t]*•\s*\d{5,}"
    r"\s+(?P<date>[A-Za-z]{3,9}\.?\s+\d{1,2},\s*\d{4})\s*•\s*"
    r"(?P<code>[A-Za-z0-9\-]+)\s*$", re.M)
ELITE_SECTION_ORDER = [("RW", 1), ("RW", 2), ("Math", 1), ("Math", 2)]


def normalize_date(date_str):
    """'AUG. 21, 2026' / '8/21/2026' -> 'August 21, 2026' (else unchanged)."""
    for fmt in ("%b. %d, %Y", "%b %d, %Y", "%B %d, %Y", "%m/%d/%Y",
                "%Y-%m-%d"):
        try:
            d = datetime.datetime.strptime(date_str.strip(), fmt)
            return f"{d.strftime('%B')} {d.day}, {d.year}"
        except ValueError:
            continue
    return date_str


def parse_elite_layout(text: str, result: StudentResult) -> bool:
    """Fill `result` from the Elite Prep score-report layout. Returns True
    when question data was recognized."""
    if "TEST SCORE REPORT" not in text.upper():
        return False
    if not ELITE_TABLE_HEADER.search(text):
        return False

    blocks, cur = [], []
    for ln in text.splitlines():
        m = ELITE_ROW.match(ln)
        if not m:
            continue
        n = int(m.group(1))
        if n == 1 and cur:
            blocks.append(cur)
            cur = []
        if not cur and n != 1:
            continue                      # stray line, not a table start
        if cur and n != cur[-1][0] + 1:
            continue                      # not sequential -> not a row
        cur.append((n, m.group(2), m.group(3), m.group(4)))
    if cur:
        blocks.append(cur)
    if not blocks:
        return False

    order = [("RW", int(a)) if a else ("Math", int(b))
             for a, b in ELITE_TABLE_HEADER.findall(text)]
    if len(order) != 4 or set(order) != set(ELITE_SECTION_ORDER):
        order = ELITE_SECTION_ORDER

    for (section, module), rows in zip(order, blocks[:4]):
        for n, correct_ans, response, code in rows:
            r = response.strip().upper().strip(".,;")
            if r == "OK":
                status = "correct"
            elif r in ("OMIT", "OMITTED", "--", "-", "BLANK"):
                status = "omitted"
            else:
                status = "incorrect"
            domain = ELITE_SKILL_CODES[section].get(code) if code else None
            result.questions.append(QuestionResult(
                section, module, n, status, domain,
                is_gridin_answer(correct_ans) if section == "Math" else None))

    # Page-1 skill summary ("CraftandStructure 86%" ... "12of14correct") -
    # kept as a fallback for reports whose rows carry no skill-code column.
    # The labels and the "N of M correct" lines are interleaved with other
    # text but appear in the same order, so they are zipped positionally.
    labels = []
    for label, pct in re.findall(r"([A-Za-z][A-Za-z\-]+)\s*(\d{1,3})%", text):
        domain = canonical_domain(re.sub(r"(?<=[a-z])(?=[A-Z])", " ", label))
        if domain:
            labels.append(domain)
    counts = re.findall(r"(\d+)\s*of\s*(\d+)\s*correct", text, re.I)
    for domain, (c, t) in zip(labels, counts):
        if domain not in result.domain_counts:
            result.domain_counts[domain] = (int(c), int(t))

    m = re.search(r"\bTotal\s*\n\s*(\d{3,4})\b", text)
    if m and 400 <= int(m.group(1)) <= 1600:
        result.total = int(m.group(1))
    m = re.search(r"Reading\s*and\s*Writing\s+(\d{3})\b", text, re.I)
    if m and 200 <= int(m.group(1)) <= 800:
        result.rw = int(m.group(1))
    m = re.search(r"^Math\s+(\d{3})\b", text, re.I | re.M)
    if m and 200 <= int(m.group(1)) <= 800:
        result.math = int(m.group(1))

    m = ELITE_ID_LINE.search(text)
    if m:
        result.student_name = m.group("name").strip()
        result.test_date = normalize_date(m.group("date"))
        result.test_code = m.group("code").strip()
    return bool(result.questions)


def parse_text(text: str, source: str) -> StudentResult:
    """Parse the extracted text of one student's score-report PDF."""
    result = StudentResult(source)

    if parse_elite_layout(text, result):
        return result

    # Scaled scores (validated against the real SAT scales).
    for key, pat in SCORE_PATTERNS.items():
        for m in pat.finditer(text):
            val = int(m.group(1))
            if key == "total" and 400 <= val <= 1600:
                result.total = val
                break
            if key in ("rw", "math") and 200 <= val <= 800:
                setattr(result, key, val)
                break

    for pat in TEST_CODE_PATTERNS:
        m = pat.search(text)
        if m:
            result.test_code = m.group(1).strip()
            break
    for pat in NAME_PATTERNS:
        m = pat.search(text)
        if m:
            cand = m.group(1).strip()
            # ignore "Name: Score Report"-style false positives
            if not re.search(r"\b(score|report|test|date|code)\b", cand, re.I):
                result.student_name = cand
            break
    for pat in DATE_PATTERNS:
        m = pat.search(text)
        if m:
            result.test_date = m.group(0)
            break

    seen = set()  # (section, module, number) - keep first occurrence
    cur_section, cur_module = None, None

    def add_question(section, module, number, status, domain, gridin):
        key = (section, module, number)
        if status and 1 <= number <= 40 and key not in seen:
            seen.add(key)
            result.questions.append(
                QuestionResult(section, module, number, status,
                               domain, gridin))

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Style A: everything on one row.
        m = ROW_A.match(line)
        if m:
            number = int(m.group(1))
            section = "Math" if m.group(2).lower().startswith("math") else "RW"
            module = int(m.group(3))
            status, domain, rest = extract_status_and_domain(m.group(4))
            ca, ya = answers_from_fragment(rest)
            if status is None:
                status = infer_status(ca, ya)
            add_question(section, module, number, status, domain,
                         is_gridin_answer(ca) if section == "Math" else None)
            continue

        # Header lines set the parsing context for style-B rows.
        header_hit = False
        if len(line) < 80:
            for pat, sec in SECTION_PATTERNS:
                if pat.search(line):
                    cur_section = sec
                    header_hit = True
                    break
            mm = MODULE_PATTERN.search(line)
            if mm:
                cur_module = int(mm.group(1))
                header_hit = True
        if header_hit and not re.match(r"^\s*\d", line):
            continue

        # Style B: bare question rows under the current section/module.
        if cur_section and cur_module:
            m = ROW_B.match(line)
            if m:
                number = int(m.group(1))
                if not (1 <= number <= 40):
                    continue
                status, domain, rest = extract_status_and_domain(m.group(2))
                ca, ya = answers_from_fragment(rest)
                if status is None:
                    status = infer_status(ca, ya)
                add_question(
                    cur_section, cur_module, number, status, domain,
                    is_gridin_answer(ca) if cur_section == "Math" else None)

    return result


def parse_pdf_stream(fileobj, name: str) -> StudentResult:
    """Parse a student's score report from a file-like object (web upload)."""
    return parse_text(extract_pdf_text(fileobj), name)


def parse_pdf(path: Path, dump_text: bool = False) -> StudentResult:
    if path.suffix.lower() == ".txt":  # debugging convenience
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        text = extract_pdf_text(str(path))

    if dump_text:
        dump_path = path.with_suffix(path.suffix + ".extracted.txt")
        dump_path.write_text(text, encoding="utf-8")
        print(f"  [debug] extracted text written to {dump_path.name}")

    return parse_text(text, path.name)


# ---------------------------------------------------------------------------
# Multi-student files (one PDF containing every student's score report)
# ---------------------------------------------------------------------------

# A new student's report starts where one of these markers re-appears.
STUDENT_BOUNDARY_PATTERNS = [
    re.compile(r"•\s*\d{5,}\s+[A-Za-z]{3,9}\.?\s+\d{1,2},\s*\d{4}"),  # Elite
    NAME_PATTERNS[0],
    re.compile(r"total\s*(?:sat\s*)?score", re.I),
    re.compile(r"score\s*report", re.I),
]


def _split_candidates(text: str):
    """Yield candidate ways of splitting one file into per-student chunks.

    Strategy 1: a boundary marker (student ID line / name / 'Total Score' /
    'Score Report') that occurs 2+ times - split before each occurrence.
    Strategy 2: page breaks (\\f) - group pages so that each group holds
    exactly one 'Total Score'.
    """
    lines = text.splitlines()
    for pat in STUDENT_BOUNDARY_PATTERNS:
        idxs = [i for i, ln in enumerate(lines) if pat.search(ln)]
        if len(idxs) >= 2:
            starts = [0] + idxs[1:]
            chunks = ["\n".join(lines[a:b])
                      for a, b in zip(starts, starts[1:] + [len(lines)])]
            yield [c for c in chunks if c.strip()]

    pages = text.split("\f")
    if len(pages) >= 2:
        chunks, cur = [], []
        for pg in pages:
            if cur and SCORE_PATTERNS["total"].search(pg) and any(
                    SCORE_PATTERNS["total"].search(c) for c in cur):
                chunks.append("\n".join(cur))
                cur = []
            cur.append(pg)
        if cur:
            chunks.append("\n".join(cur))
        if len(chunks) >= 2:
            yield chunks


def parse_text_all(text: str, source: str):
    """Parse a file that may hold one OR many students' score reports.
    Returns a list of StudentResult (only those with question data).

    A candidate split is accepted only if it yields 2+ students with
    question data and does not separate scores from questions (i.e. either
    2+ of those students carry scaled scores, or none of them do)."""
    for chunks in _split_candidates(text):
        parsed = [parse_text(c, f"{source}#{k}")
                  for k, c in enumerate(chunks, start=1)]
        ok = [s for s in parsed if s.parsed_ok]
        scored = [s for s in ok if s.total or s.rw or s.math]
        if len(ok) >= 2 and (len(scored) >= 2 or not scored):
            return ok
    s = parse_text(text, source)
    return [s] if s.parsed_ok else []


def parse_pdf_stream_all(fileobj, name: str):
    return parse_text_all(extract_pdf_text(fileobj), name)


def parse_pdf_all(path: Path, dump_text: bool = False):
    if path.suffix.lower() == ".txt":
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        text = extract_pdf_text(str(path))
    if dump_text:
        dump_path = path.with_suffix(path.suffix + ".extracted.txt")
        dump_path.write_text(text, encoding="utf-8")
        print(f"  [debug] extracted text written to {dump_path.name}")
    return parse_text_all(text, path.name)


# ---------------------------------------------------------------------------
# Question -> skill-area maps (skill_maps/<TEST CODE>.csv)
# ---------------------------------------------------------------------------
# Score reports that carry no per-question skill area (e.g. the Elite Prep
# "TEST SCORE REPORT" layout) can still get the "Average Accuracy by Skill
# Area" table when a map for the test form exists:
#
#     section,module,question,skill
#     RW,1,1,Craft and Structure
#     Math,2,17,Advanced Math
#
# section: RW / Reading and Writing / English / Math;  skill: one of the 8
# SAT domains (aliases such as "Geometry & Trigonometry" are accepted).

SKILL_MAP_DIR = Path(__file__).resolve().parent / "skill_maps"


_COMPACT_DOMAINS = {re.sub(r"[^a-z]", "", k.replace("&", "and")): v
                    for k, v in DOMAIN_ALIASES.items()}


def canonical_domain(name):
    """'Geometry & Trigonometry', 'CraftandStructure', ' algebra ' -> the
    canonical skill-area name (or None). Spaces/punctuation are ignored."""
    key = re.sub(r"[^a-z]", "", (name or "").lower().replace("&", "and"))
    return _COMPACT_DOMAINS.get(key)


def parse_skill_map_text(text: str):
    """CSV text -> {(section, module, question): domain}."""
    import csv
    import io as _io
    smap = {}
    rows = list(csv.reader(_io.StringIO(text)))
    for row in rows:
        cells = [c.strip() for c in row]
        if len(cells) < 4 or not cells[3]:
            continue
        sec, mod, q, skill = cells[:4]
        if not (mod.isdigit() and q.isdigit()):
            continue                       # header line or blank
        section = "Math" if sec.lower().startswith("m") else "RW"
        domain = canonical_domain(skill)
        if domain:
            smap[(section, int(mod), int(q))] = domain
    return smap


def load_skill_map(test_code):
    """Return the map for this test code from skill_maps/, or {}."""
    if not test_code or not SKILL_MAP_DIR.is_dir():
        return {}
    want = re.sub(r"[^a-z0-9]", "", test_code.lower())
    for p in SKILL_MAP_DIR.glob("*.csv"):
        if re.sub(r"[^a-z0-9]", "", p.stem.lower()) == want:
            return parse_skill_map_text(p.read_text(encoding="utf-8-sig"))
    return {}


def apply_skill_map(students, smap):
    """Fill in q.domain from the map where the PDF gave none. Returns the
    number of questions that received a skill area."""
    if not smap:
        return 0
    n = 0
    for s in students:
        for q in s.questions:
            if q.domain is None:
                d = smap.get((q.section, q.module, q.number))
                if d:
                    q.domain = d
                    n += 1
    return n


def parse_date_iso(date_str):
    for fmt in ("%B %d, %Y", "%m/%d/%Y", "%Y-%m-%d"):
        try:
            return datetime.datetime.strptime(date_str, fmt).date().isoformat()
        except (ValueError, TypeError):
            continue
    return None


# ---------------------------------------------------------------------------
# Class-wide aggregation
# ---------------------------------------------------------------------------

class ClassStats:
    def __init__(self, students):
        self.students = students
        self.n = len(students)

        def series(attr):
            return [getattr(s, attr) for s in students
                    if getattr(s, attr) is not None]

        self.scores = {k: series(k) for k in ("total", "rw", "math")}

        # {(section, module)}[qnum] -> count
        self.missed = defaultdict(lambda: defaultdict(int))
        self.gridin_q = defaultdict(set)      # {(sec,mod)} -> {qnum}
        self.domain_correct = defaultdict(int)
        self.domain_total = defaultdict(int)
        self.late_omissions = 0               # omitted late in a module
        self.q_domain = {}                    # (sec, mod, q) -> skill area

        for s in students:
            module_max = defaultdict(int)
            for q in s.questions:
                k = (q.section, q.module)
                module_max[k] = max(module_max[k], q.number)
            for q in s.questions:
                k = (q.section, q.module)
                if q.status != "correct":
                    self.missed[k][q.number] += 1
                if q.gridin:
                    self.gridin_q[k].add(q.number)
                if q.domain:
                    self.q_domain.setdefault(k + (q.number,), q.domain)
                    self.domain_total[q.domain] += 1
                    if q.status == "correct":
                        self.domain_correct[q.domain] += 1
                if (q.status == "omitted"
                        and q.number >= max(1, int(module_max[k] * 2 / 3))):
                    self.late_omissions += 1
            # No per-question skill areas, but a per-skill summary table.
            if (not any(q.domain for q in s.questions)
                    and s.domain_counts):
                for d, (c, t) in s.domain_counts.items():
                    self.domain_total[d] += t
                    self.domain_correct[d] += c

    def avg(self, key):
        vals = self.scores[key]
        return round(sum(vals) / len(vals)) if vals else None

    def rng(self, key):
        vals = self.scores[key]
        return (min(vals), max(vals)) if vals else None

    def domain_accuracy(self, domain):
        t = self.domain_total.get(domain, 0)
        return round(100 * self.domain_correct[domain] / t) if t else None

    def weakest_domain(self, domains):
        pairs = [(d, self.domain_accuracy(d)) for d in domains]
        pairs = [(d, a) for d, a in pairs if a is not None]
        return min(pairs, key=lambda x: x[1]) if pairs else None

    def has_domain_data(self):
        return any(self.domain_total.values())

    def is_gridin(self, section, module, qnum):
        return qnum in self.gridin_q[(section, module)]

    def error_groups(self, section, module):
        """[(missed_count, [qnums])] sorted by missed_count descending -
        i.e. most-missed questions first, as required."""
        groups = defaultdict(list)
        for qnum, c in self.missed[(section, module)].items():
            if c > 0:
                groups[c].append(qnum)
        return [(c, sorted(groups[c])) for c in sorted(groups, reverse=True)]

    def top_missed(self, section, module, min_rate=0.5, limit=3):
        """Question numbers with the highest error rates (>= min_rate)."""
        out = []
        for c, qs in self.error_groups(section, module):
            if c / self.n >= min_rate:
                out.extend(qs)
            if len(out) >= limit:
                break
        return sorted(out[:limit])


# ---------------------------------------------------------------------------
# Narrative notes (rule-based text; Claude may override via `narrative`)
# ---------------------------------------------------------------------------
# Every note is a list of (text, bold) fragments so that domain names and
# key numbers can be emphasized exactly as in the reference document.

DISPLAY_DOMAIN = {
    "Problem-Solving and Data Analysis": "Problem-Solving & Data Analysis",
    "Geometry and Trigonometry": "Geometry & Trigonometry",
}


def display_domain(d):
    return DISPLAY_DOMAIN.get(d, d)


def join_names(items):
    if len(items) <= 1:
        return items[0] if items else ""
    if len(items) == 2:
        return f"{items[0]} and {items[1]}"
    return ", ".join(items[:-1]) + f", and {items[-1]}"


def q_names(qnums):
    return join_names([f"Q{q}" for q in qnums])


def emphasize(text):
    """Plain text -> fragments with the eight skill-area names in bold."""
    names = sorted({display_domain(d) for d in RW_DOMAINS + MATH_DOMAINS}
                   | set(RW_DOMAINS + MATH_DOMAINS), key=len, reverse=True)
    pat = re.compile("|".join(re.escape(n) for n in names))
    frags, pos = [], 0
    for m in pat.finditer(text):
        if m.start() > pos:
            frags.append((text[pos:m.start()], False))
        frags.append((m.group(0), True))
        pos = m.end()
    if pos < len(text):
        frags.append((text[pos:], False))
    return frags or [(text, False)]


def overview_note(stats):
    rw_avg, math_avg = stats.avg("rw"), stats.avg("math")
    if rw_avg is None or math_avg is None:
        return None
    if math_avg >= rw_avg:
        hi, hi_v, lo, lo_v = "Math", math_avg, "Reading & Writing", rw_avg
    else:
        hi, hi_v, lo, lo_v = "Reading & Writing", rw_avg, "Math", math_avg
    verb = ("showed slightly stronger performance than"
            if abs(hi_v - lo_v) <= 30 else "showed stronger performance than")
    spread = max((r[1] - r[0]) for r in
                 (stats.rng("total"), stats.rng("rw"), stats.rng("math"))
                 if r) if stats.rng("total") else 0
    tail = ("Because the spread between students is wide, we recommend "
            "reviewing these figures alongside each student's individual "
            "Score Report." if spread >= 150 else
            "Performance was fairly consistent across the class; individual "
            "results are provided in each student's Score Report.")
    return [(f"Overall, the {hi} section (average ", False), (str(hi_v), True),
            (f") {verb} the {lo} section (average ", False), (str(lo_v), True),
            (f"). {tail}", False)]


def skill_note(stats, test_code=""):
    pairs = [(d, stats.domain_accuracy(d)) for d in RW_DOMAINS + MATH_DOMAINS]
    pairs = [(d, a) for d, a in pairs if a is not None]
    if not pairs:
        return [("Skill-area accuracy could not be computed for this test: "
                 "the score reports do not list a skill area per question "
                 "and no question-to-skill map was provided for test code "
                 f"{test_code}.", False)]
    pairs.sort(key=lambda x: x[1])
    below = [(d, a) for d, a in pairs if a < 50]
    strongest = max(pairs, key=lambda x: x[1])
    frags = []
    if below:
        frags.append(("As the table shows, ", False))
        for i, (d, a) in enumerate(below):
            if i:
                frags.append((", " if i < len(below) - 1 else " and ", False))
            frags += [(display_domain(d), True), (f" ({a}%)", False)]
        frags.append((f" fell below the 50% threshold on this administration"
                      f" and will receive focused supplementary instruction "
                      f"during class. The strongest area was ", False))
        frags += [(display_domain(strongest[0]), True),
                  (f" ({strongest[1]}%).", False)]
        return frags
    d1, a1 = pairs[0]
    frags += [("As the table shows, no single area fell below the 50% "
               "threshold on this administration. The lowest area was ",
               False), (display_domain(d1), True), (f" at {a1}%", False)]
    rest = pairs[1:3]
    if rest:
        frags.append(", followed by ")
        frags[-1] = (", followed by ", False)
        for i, (d, a) in enumerate(rest):
            if i:
                frags.append((" and ", False))
            frags += [(display_domain(d), True), (f" ({a}%)", False)]
    frags.append((". We will provide supplementary instruction focused on "
                  "these comparatively weaker areas during class.", False))
    return frags


def _top_items(stats, section, module, limit=2):
    """[(qnum, missed_count)] for the most-missed questions of a module."""
    out = []
    for c, qs in stats.error_groups(section, module):
        for q in qs:
            out.append((q, c))
            if len(out) >= limit:
                return out
    return out


def _item_list(stats, section, module, limit=2):
    items = _top_items(stats, section, module, limit)
    return join_names([f"Q{q} ({c} of {stats.n})" for q, c in items])


def math_note(stats):
    n = stats.n
    tops = {m: _top_items(stats, "Math", m) for m in (1, 2)}
    all_items = [(m, q, c) for m in (1, 2) for q, c in tops[m]]
    if not all_items:
        return [("→ In Math, no incorrect answers were recorded.", False)]
    gridins = [(m, q) for m, q, c in all_items if stats.is_gridin("Math", m, q)]
    parts = [f"Module {m} {q_names([q for q, _ in tops[m]])}"
             for m in (1, 2) if tops[m]]
    if gridins and len(gridins) == len(all_items):
        text = (f"→ In Math, every one of the {len(all_items)} most-missed "
                f"items was a student-produced response (grid-in) question "
                f"{EM} {' and '.join(parts)} {EM} pointing to a recurring "
                f"challenge with free-response setup, computation, and "
                f"answer entry.")
    else:
        detail = " and ".join(
            f"Module {m} {_item_list(stats, 'Math', m)}" for m in (1, 2)
            if tops[m])
        text = f"→ In Math, the most-missed items were {detail}"
        if gridins:
            text += (f"; {q_names(sorted(q for _, q in gridins))} "
                     f"{'were' if len(gridins) > 1 else 'was a'} "
                     f"student-produced response (grid-in) "
                     f"{'questions' if len(gridins) > 1 else 'question'}, "
                     f"so answer-entry accuracy will be reinforced.")
        else:
            text += "."
    return [(text, False)]


def rw_note(stats):
    n = stats.n
    tops = {m: _top_items(stats, "RW", m) for m in (1, 2)}
    if not (tops[1] or tops[2]):
        return [("→ In Reading & Writing, no incorrect answers were "
                 "recorded.", False)]
    # hardest single item across both modules
    best = max(((m, q, c) for m in (1, 2) for q, c in tops[m]),
               key=lambda x: x[2])
    m, q, c = best
    rate = c / n
    if rate >= 1:
        lead = f"Module {m} Q{q} was missed by every student ({c} of {n})"
    elif rate >= 0.8:
        lead = (f"Module {m} Q{q} was missed by nearly every student "
                f"({c} of {n})")
    else:
        lead = f"Module {m} Q{q} was the most-missed item ({c} of {n})"
    other = 2 if m == 1 else 1
    text = f"→ In Reading & Writing, {lead}."
    if tops[other]:
        text += (f" In Module {other}, {_item_list(stats, 'RW', other)} "
                 f"{'were the toughest items' if len(tops[other]) > 1 else 'was the toughest item'}.")
    # tie the high-error questions to the weakest R&W skill area
    weakest = stats.weakest_domain(RW_DOMAINS)
    if weakest:
        top_qs = [(mm, qq) for mm in (1, 2)
                  for qq, _ in _top_items(stats, "RW", mm, limit=4)]
        doms = [stats.q_domain.get(("RW", mm, qq)) for mm, qq in top_qs]
        hits = sum(1 for d in doms if d == weakest[0])
        if hits >= 2:
            text = text.rstrip(".") + (
                f"; several of the highest-error questions fall within the "
                f"{display_domain(weakest[0])} cluster, consistent with it "
                f"being the section's lowest skill area.")
    return emphasize(text)


RW_DRILL = {
    "Standard English Conventions": (
        "Grammar & Conventions",
        "drilling punctuation, sentence boundaries, and modifier rules"),
    "Craft and Structure": (
        "Reading Strategy",
        "training close reading of word choice, text structure, and "
        "cross-text connections"),
    "Information and Ideas": (
        "Reading Comprehension",
        "practicing evidence-based inference and data-in-text questions"),
    "Expression of Ideas": (
        "Rhetorical Synthesis",
        "practicing transition and rhetorical-synthesis questions"),
}


def plan_bullets(stats):
    """[(title, text)] for the instructional plan."""
    n = stats.n
    bullets = []
    rw_pairs = [(d, stats.domain_accuracy(d)) for d in RW_DOMAINS]
    rw_pairs = sorted([(d, a) for d, a in rw_pairs if a is not None],
                      key=lambda x: x[1])
    math_pairs = [(d, stats.domain_accuracy(d)) for d in MATH_DOMAINS]
    math_pairs = sorted([(d, a) for d, a in math_pairs if a is not None],
                        key=lambda x: x[1])

    # Reading & Writing bullets
    if rw_pairs:
        d1, a1 = rw_pairs[0]
        title, drill = RW_DRILL[d1]
        top_qs = [(mm, qq) for mm in (1, 2)
                  for qq, _ in _top_items(stats, "RW", mm, limit=4)]
        hits = sum(1 for mm, qq in top_qs
                   if stats.q_domain.get(("RW", mm, qq)) == d1)
        why = (", since several of the highest-error questions concentrated "
               "in this cluster" if hits >= 2 else "")
        bullets.append((f"Reading & Writing {EM} {title}:",
                        f"We will reinforce {display_domain(d1)} (the "
                        f"section's lowest area at {a1}%), {drill}{why}."))
        others = [display_domain(d) for d, _ in rw_pairs[1:3]]
        if others:
            bullets.append((f"Reading & Writing {EM} Reading Strategy:",
                            f"We will train {join_names(others)} strategies "
                            f"to strengthen close-reading and "
                            f"rhetorical-synthesis accuracy."))
    else:
        items = [f"Module {m} {q_names([q for q, _ in _top_items(stats, 'RW', m, 3)])}"
                 for m in (1, 2) if _top_items(stats, "RW", m)]
        bullets.append((f"Reading & Writing:",
                        "We will repeatedly drill solving strategies for the "
                        f"highest-miss items ({' and '.join(items)}) to "
                        "sharpen reading and reasoning accuracy."))

    # Math bullet
    all_items = [(m, q) for m in (1, 2) for q, _ in _top_items(stats, "Math", m)]
    gridins = [(m, q) for m, q in all_items if stats.is_gridin("Math", m, q)]
    weak_math = [display_domain(d) for d, _ in math_pairs[:2]]
    focus = (f" of {join_names(weak_math)} problems" if weak_math else "")
    if gridins and len(gridins) >= max(1, len(all_items) // 2):
        where = ("all" if len(gridins) == len(all_items)
                 else f"{len(gridins)} of the {len(all_items)}")
        bullets.append((f"Math {EM} Grid-In Questions:",
                        f"We will focus on the student-produced response "
                        f"(grid-in) items {EM} where {where} of the "
                        f"most-missed Math questions appeared {EM} "
                        f"emphasizing careful answer entry, unit checks, and "
                        f"self-verification{focus}."))
    elif math_pairs:
        d1, a1 = math_pairs[0]
        items = " and ".join(f"Module {m} {q_names([q for q, _ in _top_items(stats, 'Math', m)])}"
                             for m in (1, 2) if _top_items(stats, "Math", m))
        bullets.append((f"Math {EM} {display_domain(d1)}:",
                        f"We will reinforce {display_domain(d1)} (the "
                        f"section's lowest area at {a1}%) and review the "
                        f"highest-miss items ({items}) with full worked "
                        f"solutions in class."))
    else:
        items = " and ".join(f"Module {m} {q_names([q for q, _ in _top_items(stats, 'Math', m)])}"
                             for m in (1, 2) if _top_items(stats, "Math", m))
        bullets.append(("Math:", f"We will focus intensively on the "
                                 f"highest-miss items ({items}) and "
                                 f"reinforce the concepts behind them."))

    if stats.late_omissions > 0:
        bullets.append(("Time Management:",
                        "Because several students left later questions "
                        "unanswered (Omit), we will coach real-test pacing "
                        "so students reach and attempt every question "
                        "within each module."))
    return bullets


# ---------------------------------------------------------------------------
# DOCX building (formats measured from the reference document)
# ---------------------------------------------------------------------------

NAVY = RGBColor(0x1F, 0x38, 0x64)      # titles, labels, emphasis
BLUE = RGBColor(0x44, 0x72, 0xC4)      # heading numbers, subtitle, bullets
HEADER_BLUE = "2E5496"                 # table header fill / H2 / rules
GRAY = RGBColor(0x59, 0x59, 0x59)      # letterhead tagline, arrow notes
RED = RGBColor(0xC0, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
ZEBRA = "EEF3FB"
SUBJECT_FILL = "DCE6F5"
RED_FILL = "FCE4E4"
BORDER = "B7C4DC"


def _set(el, **attrs):
    for k, v in attrs.items():
        el.set(qn(f"w:{k}"), str(v))
    return el


def _sub(parent, tag, **attrs):
    el = OxmlElement(tag)
    _set(el, **attrs)
    parent.append(el)
    return el


def para(doc, before=None, after=None, line115=False, keep_next=False,
         bottom_border=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    if line115:
        pf.line_spacing = 1.15
    if keep_next:
        pf.keep_with_next = True
    if bottom_border:
        sz, color = bottom_border
        pPr = p._p.get_or_add_pPr()
        bdr = _sub(pPr, "w:pBdr")
        _sub(bdr, "w:bottom", val="single", sz=sz, space=0, color=color)
    return p


def run(p, text, bold=False, italic=False, size=None, color=None):
    r = p.add_run(sanitize_branding(text))
    r.bold = bold or None
    r.italic = italic or None
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def runs(p, fragments, size=None, color=None, bold_color=NAVY):
    """fragments: [(text, bold)] - bold pieces are navy, as in the reference."""
    for text, bold in fragments:
        run(p, text, bold=bold, size=size,
            color=(bold_color if bold else color))


def add_h1(doc, number, text):
    p = para(doc, before=15, after=7, keep_next=True,
             bottom_border=(12, HEADER_BLUE))
    run(p, f"{number}.  ", bold=True, size=13, color=BLUE)
    run(p, text, bold=True, size=13, color=NAVY)
    return p


def add_h2(doc, i, text):
    p = para(doc, before=11, after=5, keep_next=True)
    run(p, f"{CIRCLED[i]}  {text}", bold=True, size=11,
        color=RGBColor.from_string(HEADER_BLUE))
    return p


def add_body(doc, fragments, after=7, before=None):
    p = para(doc, before=before, after=after, line115=True)
    runs(p, fragments)
    return p


def add_arrow_note(doc, fragments, before=None, after=2):
    p = para(doc, before=before, after=after, line115=True)
    run(p, "".join(t for t, _ in fragments), italic=True, size=9.5,
        color=GRAY)   # plain italic gray, as in the reference
    return p


def make_table(doc, widths, headers, header_align):
    """Bordered table with a navy header row. widths in twips."""
    table = doc.add_table(rows=1, cols=len(widths))
    tbl = table._tbl
    tblPr = tbl.tblPr
    _set(_sub(tblPr, "w:tblW"), w=sum(widths), type="dxa")
    borders = _sub(tblPr, "w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _sub(borders, f"w:{side}", val="single", sz=4, space=0, color=BORDER)
    _set(_sub(tblPr, "w:tblLayout"), type="fixed")
    grid = tbl.tblGrid
    for gc, w in zip(grid.findall(qn("w:gridCol")), widths):
        gc.set(qn("w:w"), str(w))
    hdr = table.rows[0]
    _sub(hdr._tr.get_or_add_trPr(), "w:tblHeader")
    for cell, w, text, align in zip(hdr.cells, widths, headers, header_align):
        fill_cell(cell, w, text, fill=HEADER_BLUE, bold=True, color=WHITE,
                  align=align)
    return table


def fill_cell(cell, width, text, fill=None, bold=False, color=None,
              align="center", vmerge=None):
    cell.width = Twips(width)
    tcPr = cell._tc.get_or_add_tcPr()
    if vmerge == "restart":
        _set(_sub(tcPr, "w:vMerge"), val="restart")
    elif vmerge == "continue":
        _sub(tcPr, "w:vMerge")
    if fill:
        _set(_sub(tcPr, "w:shd"), val="clear", color="auto", fill=fill)
    mar = _sub(tcPr, "w:tcMar")
    for side, w in (("top", 60), ("left", 110), ("bottom", 60),
                    ("right", 110)):
        _set(_sub(mar, f"w:{side}"), w=w, type="dxa")
    _set(_sub(tcPr, "w:vAlign"), val="center")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run(p, text, bold=bold, size=9.5, color=color)


def add_row(table, cant_split=True):
    row = table.add_row()
    if cant_split:
        _sub(row._tr.get_or_add_trPr(), "w:cantSplit")
    return row


def add_bullet(doc, title, text):
    """'• Title: text' list item with a bold blue bullet (reference style)."""
    p = doc.add_paragraph(style="List Paragraph")
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    numPr = _sub(pPr, "w:numPr")
    _set(_sub(numPr, "w:ilvl"), val=0)
    _set(_sub(numPr, "w:numId"), val=BULLET_NUM_ID[0])
    run(p, title + " ", bold=True, color=NAVY)
    run(p, text)
    return p


BULLET_NUM_ID = [None]


def ensure_bullet_numbering(doc):
    """Create the '•' bullet definition (bold, blue, hanging indent)."""
    numbering = doc.part.numbering_part.element
    abs_ids = [int(a.get(qn("w:abstractNumId")))
               for a in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId")))
               for n in numbering.findall(qn("w:num"))]
    abs_id = max(abs_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    absnum = OxmlElement("w:abstractNum")
    _set(absnum, abstractNumId=abs_id)
    _set(_sub(absnum, "w:multiLevelType"), val="hybridMultilevel")
    lvl = _sub(absnum, "w:lvl", ilvl=0)
    _set(_sub(lvl, "w:start"), val=1)
    _set(_sub(lvl, "w:numFmt"), val="bullet")
    _set(_sub(lvl, "w:lvlText"), val="•")
    _set(_sub(lvl, "w:lvlJc"), val="left")
    pPr = _sub(lvl, "w:pPr")
    _set(_sub(pPr, "w:ind"), left=360, hanging=260)
    rPr = _sub(lvl, "w:rPr")
    _sub(rPr, "w:b")
    _set(_sub(rPr, "w:color"), val="4472C4")
    # abstractNum elements must precede num elements
    nums = numbering.findall(qn("w:num"))
    if nums:
        nums[0].addprevious(absnum)
    else:
        numbering.append(absnum)
    num = OxmlElement("w:num")
    _set(num, numId=num_id)
    _set(_sub(num, "w:abstractNumId"), val=abs_id)
    numbering.append(num)
    BULLET_NUM_ID[0] = num_id


def remove_footers(doc):
    """No footers on any page."""
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer,
                       section.even_page_footer):
            try:
                footer.is_linked_to_previous = True
                for p in list(footer.paragraphs):
                    p.clear()
            except Exception:
                pass


def _narrative_fragments(value, fallback):
    """AI text (plain string) or rule-based fragments -> fragments."""
    if isinstance(value, str) and value.strip():
        return emphasize(value.strip())
    return fallback


def build_report(stats: ClassStats, test_code: str, test_date: str,
                 out_path: Path, narrative=None):
    """narrative: optional dict of AI-written commentary (see
    ai_narrative.py) - keys 'overview', 'skill', 'math_note', 'rw_note',
    'plan' [{title, text}]. Missing keys fall back to rule-based text."""
    narrative = narrative or {}
    doc = Document()

    # --- page & base style (Letter, 0.75" margins, Calibri 10 pt) ----------
    for section in doc.sections:
        section.page_width = Twips(12240)
        section.page_height = Twips(15840)
        section.left_margin = section.right_margin = Twips(1080)
        section.top_margin = section.bottom_margin = Twips(1080)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = TEXT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.paragraph_format.space_after = Pt(0)
    ensure_bullet_numbering(doc)

    n = stats.n
    try:
        dt = datetime.datetime.strptime(test_date, "%B %d, %Y")
        month_year = dt.strftime("%B %Y")
        test_date = f"{dt.strftime('%B')} {dt.day}, {dt.year}"
    except ValueError:
        month_year = test_date

    # --- letterhead / title --------------------------------------------------
    p = para(doc, after=2)
    run(p, "ELITE PREP", bold=True, size=12, color=NAVY)
    run(p, "   |   College Admissions & Test Prep", size=9, color=GRAY)
    para(doc, bottom_border=(18, "1F3864"))
    p = para(doc, before=12, after=3)
    run(p, f"{month_year} DSAT Practice Test Results Analysis Report",
        bold=True, size=15, color=NAVY)
    p = para(doc, after=11)
    run(p, f"For Teachers   ·   Test Code {test_code}   ·   "
           f"Test Date: {test_date}", bold=True, size=9.5, color=BLUE)

    # --- intro ---------------------------------------------------------------
    p = para(doc, after=7, line115=True)
    run(p, "Dear Teachers,", bold=True, size=10.5, color=NAVY)
    add_body(doc, [(
        "This report presents an analysis of the class-wide results from "
        "the DSAT (Digital SAT) practice test administered on "
        f"{test_date}. It summarizes the overall performance of the class "
        "as a whole so that we can plan targeted review and instruction "
        "together. For each individual student's score, please refer to "
        "the personal Score Report provided separately.", False)], after=6)
    add_body(doc, [("A total of ", False),
                   (f"{n} student{'s' if n != 1 else ''}", True),
                   (" sat for this practice test. The analysis below is "
                    "presented entirely as aggregate statistics.", False)],
             after=4)

    # --- 1. Class Performance Overview ----------------------------------------
    add_h1(doc, 1, "Class Performance Overview")
    add_body(doc, [("The class averages and score ranges for this test are "
                    f"as follows. (SAT total score scale: 400{EN}1600)",
                    False)])
    widths = [3200, 3000, 3400]
    table = make_table(doc, widths,
                       ["Category", "Class Average", f"Score Range (Low{EN}High)"],
                       ["left", "center", "center"])
    for i, (label, key) in enumerate((("Total", "total"),
                                      ("Reading & Writing", "rw"),
                                      ("Math", "math"))):
        avg, rng = stats.avg(key), stats.rng(key)
        fill = ZEBRA if i % 2 == 0 else None
        cells = add_row(table).cells
        fill_cell(cells[0], widths[0], label, fill=fill, bold=True,
                  color=NAVY, align="left")
        fill_cell(cells[1], widths[1], str(avg) if avg is not None else "N/A",
                  fill=fill, bold=True)
        fill_cell(cells[2], widths[2],
                  f"{rng[0]} {EN} {rng[1]}" if rng else "N/A", fill=fill)
    note = _narrative_fragments(narrative.get("overview"), overview_note(stats))
    if note:
        add_body(doc, note, before=7, after=3)

    # --- 2. Average Accuracy by Skill Area ------------------------------------
    add_h1(doc, 2, "Average Accuracy by Skill Area")
    add_body(doc, [("Below are the class-wide average accuracy rates for the "
                    "eight detailed skill areas of the SAT. Areas with an "
                    "average accuracy below 50% are highlighted in red so "
                    "that areas requiring focused improvement can be "
                    "identified at a glance.", False)])
    widths = [2600, 4400, 2600]
    table = make_table(doc, widths, ["Subject", "Skill Area", "Average Accuracy"],
                       ["center", "left", "center"])
    for gi, (subject, domains) in enumerate((("Reading & Writing", RW_DOMAINS),
                                             ("Math", MATH_DOMAINS))):
        fill = ZEBRA if gi == 0 else None
        for j, d in enumerate(domains):
            acc = stats.domain_accuracy(d)
            cells = add_row(table).cells
            fill_cell(cells[0], widths[0], subject if j == 0 else "",
                      fill=SUBJECT_FILL, bold=True, color=NAVY,
                      vmerge="restart" if j == 0 else "continue")
            fill_cell(cells[1], widths[1], display_domain(d), fill=fill,
                      align="left")
            if acc is None:
                fill_cell(cells[2], widths[2], "N/A", fill=fill, bold=True)
            else:
                fill_cell(cells[2], widths[2], f"{acc}%", fill=fill,
                          bold=True, color=RED if acc < 50 else None)
    add_body(doc, _narrative_fragments(narrative.get("skill"),
                                       skill_note(stats, test_code)),
             before=7, after=3)

    # --- 3. Priority Review Questions by Section --------------------------------
    add_h1(doc, 3, "Priority Review Questions by Section")
    add_body(doc, [("For each section, the questions that students most "
                    "frequently answered incorrectly are listed in order of "
                    "error rate. We will prioritize these questions for "
                    "review and concept reinforcement in class. (Questions "
                    "with an error rate of 60% or higher are marked in red.)",
                    False)], after=3)
    section_order = [
        ("Math", 1, f"Math {EM} Section 1 (Module 1)"),
        ("Math", 2, f"Math {EM} Section 2 (Module 2)"),
        ("RW", 1, f"English (Reading & Writing) {EM} Section 1 (Module 1)"),
        ("RW", 2, f"English (Reading & Writing) {EM} Section 2 (Module 2)"),
    ]
    widths = [2400, 3200, 4000]
    for i, (section, module, title) in enumerate(section_order, start=1):
        add_h2(doc, i, title)
        groups = stats.error_groups(section, module)
        table = make_table(doc, widths,
                           ["Error Rate", "Students Missed", "Questions"],
                           ["center", "center", "left"])
        if not groups:
            cells = add_row(table).cells
            fill_cell(cells[0], widths[0], EM, fill=ZEBRA)
            fill_cell(cells[1], widths[1], f"0 of {n}", fill=ZEBRA)
            fill_cell(cells[2], widths[2], "No incorrect answers recorded",
                      fill=ZEBRA, align="left")
        for ri, (count, qnums) in enumerate(groups, start=1):
            rate = round(100 * count / n)
            red = rate >= 60
            fill = RED_FILL if red else (ZEBRA if ri % 2 == 1 else None)
            cells = add_row(table).cells
            fill_cell(cells[0], widths[0], f"{rate}%", fill=fill, bold=True,
                      color=RED if red else NAVY)
            fill_cell(cells[1], widths[1], f"{count} of {n}", fill=fill,
                      color=RED if red else None)
            fill_cell(cells[2], widths[2], ", ".join(f"Q{q}" for q in qnums),
                      fill=fill, bold=red, color=RED if red else None,
                      align="left")
    add_arrow_note(doc, _narrative_fragments(narrative.get("math_note"),
                                             math_note(stats)), before=6)
    add_arrow_note(doc, _narrative_fragments(narrative.get("rw_note"),
                                             rw_note(stats)))

    # --- 4. Our Instructional Plan -----------------------------------------------
    add_h1(doc, 4, "Our Instructional Plan")
    plan = narrative.get("plan")
    if not (isinstance(plan, list) and plan
            and all(isinstance(b, dict) and b.get("text") for b in plan)):
        plan = [{"title": t, "text": x} for t, x in plan_bullets(stats)]
    for b in plan:
        title = (b.get("title") or "").strip()
        if title and not title.endswith(":"):
            title += ":"
        add_bullet(doc, title, b["text"].strip())

    # --- closing (no signature line, per the rules) -------------------------------
    p = para(doc, before=9, line115=True)
    run(p, "We will continue to work alongside you in support of our "
           "students' growth. Thank you.")

    remove_footers(doc)
    scrub_student_names(doc, stats, test_code)
    doc.save(str(out_path))


def _iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer):
            yield from part.paragraphs


def scrub_student_names(doc, stats, test_code=""):
    """Privacy guard: no student name (or source file name) may appear
    anywhere in the report. Names are replaced with '[student]' and the
    result is re-checked; a leftover match raises so a bad report is never
    saved."""
    tokens = set()
    for s in stats.students:
        if s.student_name:
            tokens.add(s.student_name)
        stem = Path(s.source.split("#")[0]).stem
        if stem:
            tokens.add(stem)
            # e.g. "Score report 1603888, Arnav Kunde, DSK2402UU" -> parts
            for part in re.split(r"[,;_]+", stem):
                part = part.strip()
                if (len(part) >= 3 and not re.fullmatch(r"[\d\s\-]+", part)
                        and not re.fullmatch(r"(?i)score\s*report\s*\d*", part)
                        and part.lower() != (test_code or "").lower()):
                    tokens.add(part)
    tokens = {t for t in tokens if len(t) >= 3}
    if not tokens:
        return
    pat = re.compile("|".join(re.escape(t) for t in
                              sorted(tokens, key=len, reverse=True)), re.I)
    for p in _iter_all_paragraphs(doc):
        for run_ in p.runs:
            if pat.search(run_.text):
                run_.text = pat.sub("[student]", run_.text)
    leftover = [p.text for p in _iter_all_paragraphs(doc) if pat.search(p.text)]
    if leftover:
        raise RuntimeError("student name detected in report: "
                           + leftover[0][:80])


# ---------------------------------------------------------------------------
# Output naming
# ---------------------------------------------------------------------------

def report_filename(test_code: str) -> str:
    """'<SAT Practice Test code> SAT Test Result Analysis for Teacher.docx'"""
    safe_code = re.sub(r'[\\/:*?"<>|]', "-", test_code).strip()
    return f"{safe_code} SAT Test Result Analysis for Teacher.docx"


# ---------------------------------------------------------------------------
# Main (command line)
# ---------------------------------------------------------------------------

def main(argv=None):
    ap = argparse.ArgumentParser(
        description="Generate the Elite Prep class-wide SAT test result "
                    "analysis Word report from a folder of student "
                    "score-report PDFs.")
    ap.add_argument("folder", nargs="?", default=".",
                    help="Folder containing the student PDF score reports "
                         "(default: current folder)")
    ap.add_argument("--test-code", help="Override the auto-detected test code")
    ap.add_argument("--test-date", help="Override the auto-detected test date")
    ap.add_argument("--output-dir", help="Where to save the Word report "
                                         "(default: the PDF folder)")
    ap.add_argument("--no-ai", action="store_true",
                    help="Skip Claude Opus 4.8 and use the built-in "
                         "rule-based commentary. By default the commentary "
                         "is written by Claude Opus 4.8 when "
                         "ANTHROPIC_API_KEY is set (only anonymized class "
                         "statistics are sent).")
    ap.add_argument("--dump-text", action="store_true",
                    help="Also write each PDF's extracted raw text to a "
                         ".txt file (for debugging new report layouts)")
    args = ap.parse_args(argv)

    folder = Path(args.folder).resolve()
    if not folder.is_dir():
        sys.exit(f"ERROR: folder not found: {folder}")

    pdfs = sorted(p for p in folder.iterdir()
                  if p.suffix.lower() in (".pdf", ".txt")
                  and not p.name.endswith(".extracted.txt"))
    if not pdfs:
        sys.exit(f"ERROR: no PDF files found in {folder}")

    print(f"Found {len(pdfs)} file(s) in {folder}\n")

    students, failed = [], []
    for p in pdfs:
        print(f"Parsing {p.name} ...")
        try:
            found = parse_pdf_all(p, dump_text=args.dump_text)
        except Exception as e:
            print(f"  !! could not read file: {e}")
            failed.append(p.name)
            continue
        if found:
            if len(found) > 1:
                print(f"  {len(found)} students found in this file")
            for s in found:
                n_wrong = sum(1 for q in s.questions
                              if q.status != "correct")
                print(f"  ok: {len(s.questions)} questions ({n_wrong} "
                      f"missed), scores: total={s.total} RW={s.rw} "
                      f"Math={s.math}")
            students.extend(found)
        else:
            print("  !! no question-level data recognized "
                  "(run with --dump-text and share the .extracted.txt "
                  "file so the parser can be extended)")
            failed.append(p.name)

    if not students:
        sys.exit("\nERROR: no student data could be parsed from any file. "
                 "Re-run with --dump-text to inspect the PDF text layout.")

    test_code = (args.test_code
                 or next((s.test_code for s in students if s.test_code), None)
                 or folder.name)
    test_code = sanitize_branding(test_code).strip()

    # Question -> skill-area map: skill_maps/<code>.csv or a CSV in the folder
    smap = load_skill_map(test_code)
    for csv_path in sorted(folder.glob("*.csv")):
        smap.update(parse_skill_map_text(
            csv_path.read_text(encoding="utf-8-sig")))
    mapped = apply_skill_map(students, smap)
    if mapped:
        print(f"\nSkill-area map applied ({len(smap)} questions mapped).")
    elif not any(q.domain for s in students for q in s.questions):
        print(f"\nNote: no skill-area map found for {test_code} "
              f"(add skill_maps/{test_code}.csv to fill in section 2).")

    stats = ClassStats(students)
    test_date = (args.test_date
                 or next((s.test_date for s in students if s.test_date), None)
                 or datetime.date.today().strftime("%B %d, %Y"))

    out_dir = Path(args.output_dir).resolve() if args.output_dir else folder
    out_path = out_dir / report_filename(test_code)

    narrative = None
    if not args.no_ai and not os.environ.get("ANTHROPIC_API_KEY"):
        print("\nNote: ANTHROPIC_API_KEY is not set - using the built-in "
              "commentary. Get a key at "
              "https://console.anthropic.com/settings/keys")
    elif not args.no_ai:
        from ai_narrative import write_narrative
        print("\nAsking Claude Opus 4.8 to write the commentary ...")
        try:
            narrative = write_narrative(stats, test_code, test_date)
            print("  ok")
        except Exception as e:
            print(f"  !! AI commentary unavailable ({e}); "
                  "using the built-in rule-based text instead.")

    build_report(stats, test_code, test_date, out_path, narrative=narrative)

    print(f"\nDone. {len(students)} student(s) analyzed"
          + (f", {len(failed)} file(s) skipped: {', '.join(failed)}"
             if failed else "")
          + f"\nReport saved to:\n  {out_path}")


if __name__ == "__main__":
    main()
